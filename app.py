import os, secrets, json, asyncio, uuid, sqlite3, hashlib, base64, struct, time
from datetime import datetime, timedelta, timezone
from typing import Optional
from contextlib import contextmanager

from fastapi import FastAPI, HTTPException, WebSocket, WebSocketDisconnect, UploadFile, File, Form, Request
from fastapi.responses import HTMLResponse, FileResponse, RedirectResponse, JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from io import BytesIO

DB_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "powerapp.db")
UPLOAD_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

@contextmanager
def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode=WAL")
    conn.execute("PRAGMA foreign_keys=ON")
    try:
        yield conn
        conn.commit()
    finally:
        conn.close()

def init_db():
    with get_db() as db:
        db.executescript("""
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            uuid TEXT UNIQUE,
            username TEXT UNIQUE,
            email TEXT DEFAULT '',
            phone TEXT DEFAULT '',
            hashed_password TEXT,
            full_name TEXT DEFAULT '',
            position TEXT DEFAULT '',
            role TEXT DEFAULT 'user',
            avatar TEXT DEFAULT '',
            is_active INTEGER DEFAULT 1,
            created_at TEXT,
            updated_at TEXT
        );
        CREATE TABLE IF NOT EXISTS invite_links (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            token TEXT UNIQUE,
            created_by INTEGER REFERENCES users(id),
            role TEXT DEFAULT 'user',
            max_uses INTEGER DEFAULT 1,
            uses INTEGER DEFAULT 0,
            is_active INTEGER DEFAULT 1,
            created_at TEXT,
            expires_at TEXT
        );
        CREATE TABLE IF NOT EXISTS api_keys (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            key TEXT UNIQUE,
            name TEXT,
            user_id INTEGER REFERENCES users(id),
            is_active INTEGER DEFAULT 1,
            created_at TEXT
        );
        CREATE TABLE IF NOT EXISTS categories (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT UNIQUE,
            description TEXT
        );
        CREATE TABLE IF NOT EXISTS fixed_assets (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            account TEXT DEFAULT '',
            kfo TEXT DEFAULT '',
            kps TEXT DEFAULT '',
            employee TEXT DEFAULT '',
            storage_place TEXT DEFAULT '',
            item_number INTEGER DEFAULT 0,
            asset_name TEXT DEFAULT '',
            inventory_number TEXT DEFAULT '',
            okof TEXT DEFAULT '',
            depreciation_group TEXT DEFAULT '',
            depreciation_method TEXT DEFAULT '',
            acceptance_date TEXT DEFAULT '',
            status TEXT DEFAULT 'в эксплуатации',
            useful_life_months INTEGER DEFAULT 0,
            monthly_depreciation_rate REAL DEFAULT 0,
            depreciation_percent REAL DEFAULT 0,
            initial_cost REAL DEFAULT 0,
            residual_value REAL DEFAULT 0,
            description TEXT DEFAULT '',
            created_by INTEGER REFERENCES users(id),
            created_at TEXT,
            updated_at TEXT
        );
        CREATE TABLE IF NOT EXISTS categories_fixed (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT UNIQUE
        );
        CREATE TABLE IF NOT EXISTS warehouses (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT,
            address TEXT,
            is_active INTEGER DEFAULT 1
        );
        CREATE TABLE IF NOT EXISTS stocks (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            product_id INTEGER,
            warehouse_id INTEGER REFERENCES warehouses(id),
            quantity INTEGER DEFAULT 0,
            reserved INTEGER DEFAULT 0,
            UNIQUE(product_id, warehouse_id)
        );
        CREATE TABLE IF NOT EXISTS stock_transactions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            product_id INTEGER,
            warehouse_id INTEGER REFERENCES warehouses(id),
            to_warehouse_id INTEGER,
            transaction_type TEXT,
            quantity INTEGER,
            price REAL DEFAULT 0,
            comment TEXT,
            document_number TEXT,
            user_id INTEGER REFERENCES users(id),
            created_at TEXT
        );
        CREATE TABLE IF NOT EXISTS conversations (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT,
            is_group INTEGER DEFAULT 0,
            invite_token TEXT UNIQUE,
            avatar TEXT DEFAULT '',
            created_by INTEGER REFERENCES users(id),
            created_at TEXT
        );
        CREATE TABLE IF NOT EXISTS conversation_members (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            conversation_id INTEGER REFERENCES conversations(id),
            user_id INTEGER REFERENCES users(id),
            role TEXT DEFAULT 'member',
            is_muted INTEGER DEFAULT 0,
            joined_at TEXT
        );
        CREATE TABLE IF NOT EXISTS messages (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            conversation_id INTEGER REFERENCES conversations(id),
            user_id INTEGER REFERENCES users(id),
            content TEXT,
            encrypted_content TEXT DEFAULT '',
            msg_type TEXT DEFAULT 'text',
            file_url TEXT DEFAULT '',
            duration REAL DEFAULT 0,
            reply_to INTEGER,
            is_edited INTEGER DEFAULT 0,
            is_deleted INTEGER DEFAULT 0,
            created_at TEXT,
            updated_at TEXT
        );
        CREATE TABLE IF NOT EXISTS calls (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            conversation_id INTEGER,
            caller_id INTEGER REFERENCES users(id),
            call_type TEXT DEFAULT 'voice',
            status TEXT DEFAULT 'ringing',
            started_at TEXT,
            ended_at TEXT,
            duration_seconds REAL DEFAULT 0
        );
        CREATE TABLE IF NOT EXISTS data_transfers (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            filename TEXT,
            original_name TEXT,
            file_size INTEGER,
            mime_type TEXT,
            sender_id INTEGER REFERENCES users(id),
            receiver_id INTEGER,
            message TEXT,
            download_count INTEGER DEFAULT 0,
            is_public INTEGER DEFAULT 0,
            share_token TEXT UNIQUE,
            created_at TEXT
        );
        CREATE TABLE IF NOT EXISTS audit_log (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            action TEXT,
            details TEXT,
            ip_address TEXT,
            created_at TEXT
        );
        """)

SECRET_KEY = "powerapp-super-secret-2024-change-in-production"
ROLES_HIERARCHY = {"superadmin": 4, "admin": 3, "high_user": 2, "user": 1}

def hash_password(pw):
    salt = hashlib.sha256(SECRET_KEY.encode()).hexdigest()[:16]
    return hashlib.pbkdf2_hmac('sha256', pw.encode(), salt.encode(), 100000).hex()

def verify_password(pw, h):
    return hash_password(pw) == h

def create_token(user_id):
    exp = (datetime.now(timezone.utc) + timedelta(days=7)).isoformat()
    payload = json.dumps({"sub": user_id, "exp": exp, "iat": now_iso()})
    return base64.urlsafe_b64encode(payload.encode()).decode()

def decode_token(token):
    try:
        return json.loads(base64.urlsafe_b64decode(token.encode()))
    except Exception:
        return None

def create_api_key():
    return f"pa_{secrets.token_urlsafe(32)}"

def now_iso():
    return datetime.now(timezone.utc).isoformat()

def row_to_dict(row):
    return dict(row) if row else None

def rows_to_list(rows):
    return [dict(r) for r in rows]

def safe_user(u):
    if not u: return None
    d = dict(u)
    d.pop("hashed_password", None)
    return d

def get_user(request: Request):
    token = request.cookies.get("token")
    auth = request.headers.get("Authorization", "")
    api_key = request.headers.get("X-API-Key")
    with get_db() as db:
        if api_key:
            k = db.execute("SELECT * FROM api_keys WHERE key=? AND is_active=1", (api_key,)).fetchone()
            if k:
                return safe_user(db.execute("SELECT * FROM users WHERE id=? AND is_active=1", (k["user_id"],)).fetchone())
        if auth.startswith("Bearer "):
            payload = decode_token(auth[7:])
            if payload:
                return safe_user(db.execute("SELECT * FROM users WHERE id=? AND is_active=1", (payload["sub"],)).fetchone())
        if token:
            payload = decode_token(token)
            if payload:
                return safe_user(db.execute("SELECT * FROM users WHERE id=? AND is_active=1", (payload["sub"],)).fetchone())
    raise HTTPException(401, "Not authenticated")

def require_role(min_role):
    def check(request: Request):
        u = get_user(request)
        if ROLES_HIERARCHY.get(u["role"], 0) < ROLES_HIERARCHY.get(min_role, 0):
            raise HTTPException(403, "Insufficient permissions")
        return u
    return check

def log_action(user_id, action, details="", ip=""):
    with get_db() as db:
        db.execute("INSERT INTO audit_log (user_id, action, details, ip_address, created_at) VALUES (?,?,?,?,?)",
                   (user_id, action, details, ip, now_iso()))

ws_connections: dict[int, list[WebSocket]] = {}
call_signals: dict[int, list[WebSocket]] = {}

app = FastAPI(title="PowerApp")
app.mount("/static", StaticFiles(directory=os.path.join(os.path.dirname(os.path.abspath(__file__)), "static")), name="static")
templates = Jinja2Templates(directory=os.path.join(os.path.dirname(os.path.abspath(__file__)), "templates"))

@app.on_event("startup")
def startup():
    init_db()
    with get_db() as db:
        if not db.execute("SELECT 1 FROM users LIMIT 1").fetchone():
            db.execute("""INSERT INTO users (uuid, username, email, phone, hashed_password, full_name, position, role, created_at)
                         VALUES (?,?,?,?,?,?,?,?,?)""",
                       (str(uuid.uuid4()), "superadmin", "super@powerapp.local", "+70000000000",
                        hash_password("superadmin"), "Суперадминистратор", "Главный системный администратор", "superadmin", now_iso()))
            db.execute("INSERT INTO categories_fixed (name) VALUES (?)", ("Здания",))
            db.execute("INSERT INTO categories_fixed (name) VALUES (?)", ("Машины и оборудование",))
            db.execute("INSERT INTO warehouses (name, address) VALUES (?,?)", ("Основной склад", "Главный"))

def log_audit(user_id, action, details="", request=None):
    ip = ""
    if request:
        ip = request.client.host if request.client else ""
    log_action(user_id, action, details, ip)

# ═══════════════════════════════════════════
# AUTH
# ═══════════════════════════════════════════
@app.post("/api/auth/register-by-link")
def register_by_link(data: dict):
    with get_db() as db:
        inv = db.execute("SELECT * FROM invite_links WHERE token=? AND is_active=1", (data["token"],)).fetchone()
        if not inv: raise HTTPException(400, "Invalid or expired invite link")
        if inv["max_uses"] > 0 and inv["uses"] >= inv["max_uses"]: raise HTTPException(400, "Invite link exhausted")
        if inv["expires_at"] and inv["expires_at"] < now_iso(): raise HTTPException(400, "Invite link expired")
        db.execute("UPDATE invite_links SET uses=uses+1 WHERE id=?", (inv["id"],))
        role = inv["role"] or "user"
        uid = str(uuid.uuid4())
        db.execute("""INSERT INTO users (uuid, username, phone, hashed_password, full_name, position, role, created_at)
                     VALUES (?,?,?,?,?,?,?,?)""",
                   (uid, data.get("username", uid[:8]), data.get("phone",""), hash_password(data["password"]),
                    data["full_name"], data.get("position",""), role, now_iso()))
        u = db.execute("SELECT * FROM users WHERE uuid=?", (uid,)).fetchone()
        log_audit(u["id"], "register_by_link", f"invite={data['token']}")
        token = create_token(u["id"])
        resp = JSONResponse({"token": token, "user": safe_user(u)})
        resp.set_cookie("token", token, max_age=7*86400)
        return resp

@app.post("/api/auth/register-by-id")
def register_by_id(data: dict):
    with get_db() as db:
        inv = db.execute("SELECT * FROM invite_links WHERE token=? AND is_active=1", (data["token"],)).fetchone()
        if not inv: raise HTTPException(400, "Invalid invite link")
        if inv["max_uses"] > 0 and inv["uses"] >= inv["max_uses"]: raise HTTPException(400, "Link exhausted")
        db.execute("UPDATE invite_links SET uses=uses+1 WHERE id=?", (inv["id"],))
        role = inv["role"] or "user"
        uid = data.get("custom_id", str(uuid.uuid4()))
        db.execute("""INSERT INTO users (uuid, username, hashed_password, full_name, position, role, created_at)
                     VALUES (?,?,?,?,?,?,?)""",
                   (uid, uid, hash_password(data["password"]), data["full_name"], data.get("position",""), role, now_iso()))
        u = db.execute("SELECT * FROM users WHERE uuid=?", (uid,)).fetchone()
        token = create_token(u["id"])
        resp = JSONResponse({"token": token, "user": safe_user(u)})
        resp.set_cookie("token", token, max_age=7*86400)
        return resp

@app.post("/api/auth/admin/create-user")
def admin_create_user(data: dict, request: Request):
    u = get_user(request)
    if ROLES_HIERARCHY.get(u["role"], 0) < ROLES_HIERARCHY["admin"]:
        raise HTTPException(403, "Only admin+ can create users")
    with get_db() as db:
        if db.execute("SELECT 1 FROM users WHERE username=?", (data["username"],)).fetchone():
            raise HTTPException(400, "Username exists")
        target_role = data.get("role", "user")
        if ROLES_HIERARCHY.get(target_role, 0) >= ROLES_HIERARCHY.get(u["role"], 0):
            raise HTTPException(403, "Cannot create user with equal or higher role")
        uid = str(uuid.uuid4())
        pw = data.get("password", secrets.token_urlsafe(8))
        db.execute("""INSERT INTO users (uuid, username, phone, email, hashed_password, full_name, position, role, created_at)
                     VALUES (?,?,?,?,?,?,?,?,?)""",
                   (uid, data["username"], data.get("phone",""), data.get("email",""),
                    hash_password(pw), data["full_name"], data.get("position",""), target_role, now_iso()))
        log_audit(u["id"], "admin_create_user", f"created={data['username']} role={target_role}", request)
        return {"ok": True, "temp_password": pw, "uuid": uid}

@app.post("/api/auth/admin/create-invite")
def create_invite(data: dict, request: Request):
    u = get_user(request)
    if ROLES_HIERARCHY.get(u["role"], 0) < ROLES_HIERARCHY["admin"]:
        raise HTTPException(403, "Only admin+ can create invite links")
    with get_db() as db:
        token = secrets.token_urlsafe(16)
        target_role = data.get("role", "user")
        if ROLES_HIERARCHY.get(target_role, 0) >= ROLES_HIERARCHY.get(u["role"], 0):
            raise HTTPException(403, "Cannot invite with equal or higher role")
        expires = None
        if data.get("expires_hours"):
            expires = (datetime.now(timezone.utc) + timedelta(hours=int(data["expires_hours"]))).isoformat()
        db.execute("INSERT INTO invite_links (token, created_by, role, max_uses, expires_at, created_at) VALUES (?,?,?,?,?,?)",
                   (token, u["id"], target_role, int(data.get("max_uses", 10)), expires, now_iso()))
        log_audit(u["id"], "create_invite", f"token={token} role={target_role}", request)
        return {"token": token, "link": f"{request.base_url}invite/{token}", "role": target_role}

@app.post("/api/auth/login")
def login(data: dict):
    with get_db() as db:
        u = db.execute("SELECT * FROM users WHERE (username=? OR phone=? OR email=?) AND is_active=1",
                       (data.get("username",""), data.get("username",""), data.get("username",""))).fetchone()
        if not u or not verify_password(data["password"], u["hashed_password"]):
            raise HTTPException(401, "Invalid credentials")
        token = create_token(u["id"])
        log_audit(u["id"], "login", request=None)
        resp = JSONResponse({"token": token, "user": safe_user(u)})
        resp.set_cookie("token", token, max_age=7*86400)
        return resp

@app.get("/api/auth/me")
def get_me(request: Request):
    return get_user(request)

@app.put("/api/auth/profile")
def update_profile(data: dict, request: Request):
    u = get_user(request)
    with get_db() as db:
        db.execute("UPDATE users SET full_name=?, position=?, phone=?, email=?, updated_at=? WHERE id=?",
                   (data.get("full_name", u["full_name"]), data.get("position", u["position"]),
                    data.get("phone", u["phone"]), data.get("email", u["email"]), now_iso(), u["id"]))
    return {"ok": True}

@app.get("/api/users")
def list_users(request: Request):
    u = get_user(request)
    with get_db() as db:
        users = db.execute("SELECT * FROM users WHERE is_active=1 ORDER BY role, full_name").fetchall()
        result = []
        for usr in users:
            d = safe_user(usr)
            if ROLES_HIERARCHY.get(u["role"], 0) < ROLES_HIERARCHY["admin"] and d["role"] == "superadmin":
                continue
            result.append(d)
        return result

@app.get("/api/users/{uid}")
def get_user_profile(uid: int, request: Request):
    get_user(request)
    with get_db() as db:
        u = db.execute("SELECT id, uuid, username, full_name, position, phone, email, role, created_at FROM users WHERE id=?", (uid,)).fetchone()
        if not u: raise HTTPException(404, "User not found")
        return dict(u)

@app.put("/api/users/{uid}/role")
def change_role(uid: int, data: dict, request: Request):
    u = get_user(request)
    if u["role"] != "superadmin": raise HTTPException(403, "Only superadmin can change roles")
    with get_db() as db:
        target = db.execute("SELECT * FROM users WHERE id=?", (uid,)).fetchone()
        if not target: raise HTTPException(404, "Not found")
        new_role = data["role"]
        if new_role == "superadmin" and u["id"] != uid:
            raise HTTPException(403, "Cannot make others superadmin")
        db.execute("UPDATE users SET role=? WHERE id=?", (new_role, uid))
        log_audit(u["id"], "change_role", f"target={uid} role={new_role}", request)
    return {"ok": True}

@app.delete("/api/users/{uid}")
def deactivate_user(uid: int, request: Request):
    u = get_user(request)
    if ROLES_HIERARCHY.get(u["role"], 0) < ROLES_HIERARCHY["admin"]:
        raise HTTPException(403, "Only admin+ can deactivate users")
    with get_db() as db:
        target = db.execute("SELECT * FROM users WHERE id=?", (uid,)).fetchone()
        if not target: raise HTTPException(404, "Not found")
        if target["role"] == "superadmin": raise HTTPException(403, "Cannot deactivate superadmin")
        if ROLES_HIERARCHY.get(target["role"], 0) >= ROLES_HIERARCHY.get(u["role"], 0):
            raise HTTPException(403, "Cannot deactivate user with equal or higher role")
        db.execute("UPDATE users SET is_active=0 WHERE id=?", (uid,))
        log_audit(u["id"], "deactivate_user", f"target={uid}", request)
    return {"ok": True}

@app.post("/api/auth/api-keys")
def create_key(data: dict, request: Request):
    u = get_user(request)
    with get_db() as db:
        key = create_api_key()
        db.execute("INSERT INTO api_keys (key, name, user_id, created_at) VALUES (?,?,?,?)", (key, data["name"], u["id"], now_iso()))
        return {"key": key, "name": data["name"]}

@app.get("/api/auth/api-keys")
def list_keys(request: Request):
    u = get_user(request)
    with get_db() as db:
        keys = db.execute("SELECT * FROM api_keys WHERE user_id=?", (u["id"],)).fetchall()
        return [{"id": k["id"], "key": k["key"][:12]+"...", "name": k["name"], "created_at": k["created_at"]} for k in keys]

@app.delete("/api/auth/api-keys/{kid}")
def delete_key(kid: int, request: Request):
    u = get_user(request)
    with get_db() as db:
        db.execute("DELETE FROM api_keys WHERE id=? AND user_id=?", (kid, u["id"]))
    return {"ok": True}

@app.get("/api/audit")
def get_audit_log(request: Request):
    u = get_user(request)
    if ROLES_HIERARCHY.get(u["role"], 0) < ROLES_HIERARCHY["admin"]:
        raise HTTPException(403, "Admin+ only")
    with get_db() as db:
        logs = db.execute("""SELECT a.*, u.username FROM audit_log a
                           LEFT JOIN users u ON a.user_id=u.id ORDER BY a.created_at DESC LIMIT 200""").fetchall()
        return rows_to_list(logs)

# ═══════════════════════════════════════════
# FIXED ASSETS (ОСНОВНЫЕ СРЕДСТВА)
# ═══════════════════════════════════════════
@app.get("/api/assets")
def list_assets(request: Request):
    get_user(request)
    with get_db() as db:
        return rows_to_list(db.execute("SELECT * FROM fixed_assets ORDER BY item_number").fetchall())

@app.post("/api/assets")
def create_asset(data: dict, request: Request):
    u = get_user(request)
    with get_db() as db:
        db.execute("""INSERT INTO fixed_assets (account,kfo,kps,employee,storage_place,item_number,asset_name,
                     inventory_number,okof,depreciation_group,depreciation_method,acceptance_date,status,
                     useful_life_months,monthly_depreciation_rate,depreciation_percent,initial_cost,residual_value,
                     description,created_by,created_at,updated_at) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                   (data.get("account",""), data.get("kfo",""), data.get("kps",""), data.get("employee",""),
                    data.get("storage_place",""), int(data.get("item_number",0)), data.get("asset_name",""),
                    data.get("inventory_number",""), data.get("okof",""), data.get("depreciation_group",""),
                    data.get("depreciation_method",""), data.get("acceptance_date",""),
                    data.get("status","в эксплуатации"), int(data.get("useful_life_months",0)),
                    float(data.get("monthly_depreciation_rate",0)), float(data.get("depreciation_percent",0)),
                    float(data.get("initial_cost",0)), float(data.get("residual_value",0)),
                    data.get("description",""), u["id"], now_iso(), now_iso()))
        log_audit(u["id"], "create_asset", f"name={data.get('asset_name','')}", request)
        return {"ok": True}

@app.put("/api/assets/{aid}")
def update_asset(aid: int, data: dict, request: Request):
    u = get_user(request)
    if ROLES_HIERARCHY.get(u["role"], 0) < ROLES_HIERARCHY["high_user"]:
        raise HTTPException(403, "high_user+ required")
    with get_db() as db:
        fields = ["account","kfo","kps","employee","storage_place","item_number","asset_name",
                  "inventory_number","okof","depreciation_group","depreciation_method","acceptance_date",
                  "status","useful_life_months","monthly_depreciation_rate","depreciation_percent",
                  "initial_cost","residual_value","description"]
        sets, vals = [], []
        for f in fields:
            if f in data:
                sets.append(f"{f}=?")
                vals.append(data[f])
        if sets:
            sets.append("updated_at=?")
            vals.append(now_iso())
            vals.append(aid)
            db.execute(f"UPDATE fixed_assets SET {','.join(sets)} WHERE id=?", vals)
        log_audit(u["id"], "update_asset", f"id={aid}", request)
        return {"ok": True}

@app.delete("/api/assets/{aid}")
def delete_asset(aid: int, request: Request):
    u = get_user(request)
    if ROLES_HIERARCHY.get(u["role"], 0) < ROLES_HIERARCHY["admin"]:
        raise HTTPException(403, "Admin+ required")
    with get_db() as db:
        db.execute("DELETE FROM fixed_assets WHERE id=?", (aid,))
        log_audit(u["id"], "delete_asset", f"id={aid}", request)
    return {"ok": True}

@app.get("/api/assets/export/{fmt}")
def export_assets(fmt: str, request: Request):
    u = get_user(request)
    with get_db() as db:
        assets = rows_to_list(db.execute("SELECT * FROM fixed_assets ORDER BY item_number").fetchall())
    if fmt == "csv":
        headers = ["№","Счет","КФО","КПС","Сотрудник","Место хранения","Наименование","Инв.номер",
                   "ОКОФ","Грп.аморт.","Способ аморт.","Дата принятия","Состояние","Срок исп.(мес.)",
                   "Норма износа %","Износ %","Нач.стоимость","Ост.стоимость"]
        lines = [",".join(headers)]
        for a in assets:
            line = [str(a.get(h.lower().replace(" ","_").replace(".","").replace("№","item_number").replace("%","").replace("(","").replace(")",""), "")) for h in headers]
            lines.append(";".join([
                str(a.get("item_number","")), a.get("account",""), a.get("kfo",""), a.get("kps",""),
                a.get("employee",""), a.get("storage_place",""), a.get("asset_name",""),
                a.get("inventory_number",""), a.get("okof",""), a.get("depreciation_group",""),
                a.get("depreciation_method",""), a.get("acceptance_date",""), a.get("status",""),
                str(a.get("useful_life_months","")), str(a.get("monthly_depreciation_rate","")),
                str(a.get("depreciation_percent","")), str(a.get("initial_cost","")),
                str(a.get("residual_value",""))
            ]))
        content = "\n".join(lines)
        return StreamingResponse(BytesIO(content.encode("utf-8-sig")), media_type="text/csv",
                               headers={"Content-Disposition": f"attachment; filename=assets_{now_iso()[:10]}.csv"})

    elif fmt == "xlsx":
        try:
            import openpyxl
        except ImportError:
            try:
                import subprocess
                subprocess.check_call(["pip", "install", "openpyxl", "--only-binary=:all:"])
                import openpyxl
            except:
                raise HTTPException(500, "Excel export not available")
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Основные средства"
        headers = ["№ п/п","Счет","КФО","КПС","ЦМО.Сотрудник","ЦМО.Место хранения","Основное средство",
                   "Инв. номер","ОКОФ","Аморт. группа","Способ начисления","Дата принятия","Состояние",
                   "Срок полезн. исп.","Норма износа %","Износ %","Начальная стоимость","Остаточная стоимость"]
        for col, h in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=h)
            cell.font = openpyxl.styles.Font(bold=True)
            cell.fill = openpyxl.styles.PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
            cell.font = openpyxl.styles.Font(bold=True, color="FFFFFF")
        for row, a in enumerate(assets, 2):
            ws.cell(row=row, column=1, value=a.get("item_number"))
            ws.cell(row=row, column=2, value=a.get("account"))
            ws.cell(row=row, column=3, value=a.get("kfo"))
            ws.cell(row=row, column=4, value=a.get("kps"))
            ws.cell(row=row, column=5, value=a.get("employee"))
            ws.cell(row=row, column=6, value=a.get("storage_place"))
            ws.cell(row=row, column=7, value=a.get("asset_name"))
            ws.cell(row=row, column=8, value=a.get("inventory_number"))
            ws.cell(row=row, column=9, value=a.get("okof"))
            ws.cell(row=row, column=10, value=a.get("depreciation_group"))
            ws.cell(row=row, column=11, value=a.get("depreciation_method"))
            ws.cell(row=row, column=12, value=a.get("acceptance_date"))
            ws.cell(row=row, column=13, value=a.get("status"))
            ws.cell(row=row, column=14, value=a.get("useful_life_months"))
            ws.cell(row=row, column=15, value=a.get("monthly_depreciation_rate"))
            ws.cell(row=row, column=16, value=a.get("depreciation_percent"))
            ws.cell(row=row, column=17, value=a.get("initial_cost"))
            ws.cell(row=row, column=18, value=a.get("residual_value"))
        for col in ws.columns:
            ws.column_dimensions[col[0].column_letter].width = 18
        buf = BytesIO()
        wb.save(buf)
        buf.seek(0)
        return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                               headers={"Content-Disposition": f"attachment; filename=assets_{now_iso()[:10]}.xlsx"})

    elif fmt == "docx":
        try:
            from docx import Document
        except ImportError:
            try:
                import subprocess
                subprocess.check_call(["pip", "install", "python-docx", "--only-binary=:all:"])
                from docx import Document
            except:
                raise HTTPException(500, "Word export not available")
        doc = Document()
        doc.add_heading("Основные средства — Реестр", 0)
        doc.add_paragraph(f"Дата: {now_iso()[:10]}")
        table = doc.add_table(rows=1, cols=8)
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        for i, h in enumerate(["№","Наименование","Инв.номер","ОКОФ","Состояние","Нач.стоимость","Износ %","Ост.стоимость"]):
            hdr[i].text = h
        for a in assets:
            row = table.add_row().cells
            row[0].text = str(a.get("item_number",""))
            row[1].text = a.get("asset_name","")
            row[2].text = a.get("inventory_number","")
            row[3].text = a.get("okof","")
            row[4].text = a.get("status","")
            row[5].text = str(a.get("initial_cost",""))
            row[6].text = str(a.get("depreciation_percent",""))
            row[7].text = str(a.get("residual_value",""))
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                               headers={"Content-Disposition": f"attachment; filename=assets_{now_iso()[:10]}.docx"})

    elif fmt == "1c":
        lines = ["1С-ЭКСПОРТ|Основные средства|" + now_iso()[:10]]
        for a in assets:
            lines.append("|".join(["ОС", str(a.get("item_number","")), a.get("asset_name",""),
                                   a.get("inventory_number",""), a.get("okof",""), a.get("account",""),
                                   str(a.get("initial_cost","")), str(a.get("residual_value","")),
                                   a.get("status",""), a.get("depreciation_method",""),
                                   str(a.get("useful_life_months","")), a.get("acceptance_date","")]))
        content = "\n".join(lines)
        return StreamingResponse(BytesIO(content.encode("utf-8")), media_type="text/plain",
                               headers={"Content-Disposition": f"attachment; filename=assets_1c_{now_iso()[:10]}.txt"})

# ═══════════════════════════════════════════
# WAREHOUSE
# ═══════════════════════════════════════════
@app.get("/api/warehouse/summary")
def summary():
    with get_db() as db:
        ta = db.execute("SELECT COUNT(*) c FROM fixed_assets").fetchone()["c"]
        tp = db.execute("SELECT COUNT(*) c FROM stocks").fetchone()["c"]
        tw = db.execute("SELECT COUNT(*) c FROM warehouses WHERE is_active=1").fetchone()["c"]
        tv = db.execute("SELECT COALESCE(SUM(initial_cost),0) c FROM fixed_assets").fetchone()["c"]
        return {"total_assets": ta, "total_stock": tp, "total_warehouses": tw, "total_value": float(tv)}

@app.get("/api/warehouse/alerts")
def stock_alerts():
    with get_db() as db:
        rows = db.execute("SELECT s.*, a.asset_name, a.inventory_number, w.name as warehouse_name FROM stocks s LEFT JOIN fixed_assets a ON s.product_id=a.id JOIN warehouses w ON s.warehouse_id=w.id WHERE s.quantity <= 0").fetchall()
        return rows_to_list(rows)

@app.get("/api/warehouse/warehouses")
def list_warehouses():
    with get_db() as db:
        return rows_to_list(db.execute("SELECT * FROM warehouses WHERE is_active=1").fetchall())

@app.post("/api/warehouse/warehouses")
def create_warehouse(data: dict, request: Request):
    u = get_user(request)
    if ROLES_HIERARCHY.get(u["role"], 0) < ROLES_HIERARCHY["high_user"]: raise HTTPException(403)
    with get_db() as db:
        db.execute("INSERT INTO warehouses (name, address) VALUES (?,?)", (data["name"], data.get("address")))
    return {"ok": True}

@app.get("/api/warehouse/transactions")
def list_transactions(limit: int = 100):
    with get_db() as db:
        return rows_to_list(db.execute("SELECT t.*, w.name as warehouse_name FROM stock_transactions t JOIN warehouses w ON t.warehouse_id=w.id ORDER BY t.created_at DESC LIMIT ?", (limit,)).fetchall())

@app.post("/api/warehouse/transactions")
def create_transaction(data: dict, request: Request):
    u = get_user(request)
    with get_db() as db:
        db.execute("INSERT INTO stock_transactions (product_id,warehouse_id,to_warehouse_id,transaction_type,quantity,price,comment,document_number,user_id,created_at) VALUES (?,?,?,?,?,?,?,?,?,?)",
                   (data.get("product_id"), data["warehouse_id"], data.get("to_warehouse_id"), data["transaction_type"],
                    int(data["quantity"]), float(data.get("price",0)), data.get("comment"), data.get("document_number"), u["id"], now_iso()))
    return {"ok": True}

# ═══════════════════════════════════════════
# CHAT
# ═══════════════════════════════════════════
@app.get("/api/chat/conversations")
def list_conversations(request: Request):
    u = get_user(request)
    with get_db() as db:
        mids = [r["conversation_id"] for r in db.execute("SELECT conversation_id FROM conversation_members WHERE user_id=?", (u["id"],)).fetchall()]
        result = []
        for cid in mids:
            c = db.execute("SELECT * FROM conversations WHERE id=?", (cid,)).fetchone()
            if not c: continue
            mc = db.execute("SELECT COUNT(*) c FROM conversation_members WHERE conversation_id=?", (cid,)).fetchone()["c"]
            last = db.execute("SELECT m.content, u.full_name, u.username FROM messages m JOIN users u ON m.user_id=u.id WHERE m.conversation_id=? ORDER BY m.created_at DESC LIMIT 1", (cid,)).fetchone()
            members = db.execute("""SELECT cm.user_id, u.full_name, u.username, u.position FROM conversation_members cm
                                  JOIN users u ON cm.user_id=u.id WHERE cm.conversation_id=?""", (cid,)).fetchall()
            result.append({"id": c["id"], "title": c["title"], "is_group": c["is_group"], "invite_token": c["invite_token"],
                          "created_by": c["created_by"], "member_count": mc,
                          "last_message": last["content"] if last else None,
                          "last_sender": last["full_name"] or last["username"] if last else None,
                          "members": [dict(m) for m in members]})
        return result

@app.post("/api/chat/conversations")
def create_conversation(data: dict, request: Request):
    u = get_user(request)
    with get_db() as db:
        inv = secrets.token_urlsafe(16)
        db.execute("INSERT INTO conversations (title, is_group, invite_token, created_by, created_at) VALUES (?,?,?,?,?)",
                   (data["title"], int(data.get("is_group", False)), inv, u["id"], now_iso()))
        cid = db.execute("SELECT last_insert_rowid()").fetchone()[0]
        db.execute("INSERT INTO conversation_members (conversation_id, user_id, role, joined_at) VALUES (?,?,?,?)", (cid, u["id"], "admin", now_iso()))
        for mid in data.get("member_ids", []):
            if mid != u["id"]:
                db.execute("INSERT INTO conversation_members (conversation_id, user_id, joined_at) VALUES (?,?,?)", (cid, mid, now_iso()))
        return {"id": cid, "title": data["title"], "invite_token": inv}

@app.post("/api/chat/join/{token}")
def join_chat(token: str, request: Request):
    u = get_user(request)
    with get_db() as db:
        c = db.execute("SELECT * FROM conversations WHERE invite_token=?", (token,)).fetchone()
        if not c: raise HTTPException(404, "Chat not found")
        existing = db.execute("SELECT 1 FROM conversation_members WHERE conversation_id=? AND user_id=?", (c["id"], u["id"])).fetchone()
        if existing: return {"id": c["id"], "title": c["title"], "already_member": True}
        db.execute("INSERT INTO conversation_members (conversation_id, user_id, joined_at) VALUES (?,?,?)", (c["id"], u["id"], now_iso()))
        return {"id": c["id"], "title": c["title"], "already_member": False}

@app.post("/api/chat/conversations/{cid}/members")
def add_members(cid: int, data: dict, request: Request):
    u = get_user(request)
    with get_db() as db:
        member = db.execute("SELECT * FROM conversation_members WHERE conversation_id=? AND user_id=?", (cid, u["id"])).fetchone()
        if not member: raise HTTPException(403, "Not a member")
        added = 0
        for mid in data.get("user_ids", []):
            exists = db.execute("SELECT 1 FROM conversation_members WHERE conversation_id=? AND user_id=?", (cid, mid)).fetchone()
            if not exists:
                db.execute("INSERT INTO conversation_members (conversation_id, user_id, joined_at) VALUES (?,?,?)", (cid, mid, now_iso()))
                added += 1
        return {"added": added}

@app.get("/api/chat/conversations/{cid}/messages")
def get_messages(cid: int, skip: int = 0, limit: int = 100, request: Request = None):
    u = get_user(request)
    with get_db() as db:
        if not db.execute("SELECT 1 FROM conversation_members WHERE conversation_id=? AND user_id=?", (cid, u["id"])).fetchone():
            raise HTTPException(403, "Not a member")
        msgs = db.execute("""SELECT m.*, u.username, u.full_name, u.position, u.role as user_role FROM messages m
                            JOIN users u ON m.user_id=u.id WHERE m.conversation_id=? AND m.is_deleted=0
                            ORDER BY m.created_at DESC LIMIT ? OFFSET ?""",
                         (cid, limit, skip)).fetchall()
        return rows_to_list(list(reversed(msgs)))

@app.post("/api/chat/conversations/{cid}/messages")
def send_message(cid: int, data: dict, request: Request):
    u = get_user(request)
    with get_db() as db:
        if not db.execute("SELECT 1 FROM conversation_members WHERE conversation_id=? AND user_id=?", (cid, u["id"])).fetchone():
            raise HTTPException(403, "Not a member")
        msg_type = data.get("msg_type", "text")
        db.execute("""INSERT INTO messages (conversation_id, user_id, content, encrypted_content, msg_type, file_url, duration, reply_to, created_at, updated_at)
                     VALUES (?,?,?,?,?,?,?,?,?,?)""",
                   (cid, u["id"], data["content"], data.get("encrypted_content",""), msg_type,
                    data.get("file_url",""), float(data.get("duration",0)), data.get("reply_to"),
                    now_iso(), now_iso()))
        mid = db.execute("SELECT last_insert_rowid()").fetchone()[0]
        msg = {"type": "message", "conversation_id": cid, "user_id": u["id"], "username": u["username"],
               "full_name": u.get("full_name",""), "position": u.get("position",""), "role": u.get("role",""),
               "content": data["content"], "encrypted_content": data.get("encrypted_content",""),
               "msg_type": msg_type, "file_url": data.get("file_url",""), "duration": float(data.get("duration",0)),
               "reply_to": data.get("reply_to"), "created_at": now_iso(), "message_id": mid}
        if cid in ws_connections:
            for ws in ws_connections[cid]:
                try: asyncio.get_event_loop().create_task(ws.send_text(json.dumps(msg)))
                except: pass
        return {"id": mid}

@app.put("/api/chat/messages/{mid}")
def edit_message(mid: int, data: dict, request: Request):
    u = get_user(request)
    with get_db() as db:
        msg = db.execute("SELECT * FROM messages WHERE id=? AND user_id=?", (mid, u["id"])).fetchone()
        if not msg: raise HTTPException(404, "Not found or not yours")
        db.execute("UPDATE messages SET content=?, is_edited=1, updated_at=? WHERE id=?", (data["content"], now_iso(), mid))
    return {"ok": True}

@app.delete("/api/chat/messages/{mid}")
def delete_message(mid: int, request: Request):
    u = get_user(request)
    with get_db() as db:
        msg = db.execute("SELECT * FROM messages WHERE id=?", (mid,)).fetchone()
        if not msg: raise HTTPException(404, "Not found")
        if msg["user_id"] != u["id"] and ROLES_HIERARCHY.get(u["role"], 0) < ROLES_HIERARCHY["admin"]:
            raise HTTPException(403, "No permission")
        db.execute("UPDATE messages SET is_deleted=1, content='[Удалено]', updated_at=? WHERE id=?", (now_iso(), mid))
    return {"ok": True}

@app.websocket("/api/chat/ws/{cid}")
async def chat_ws(websocket: WebSocket, cid: int):
    await websocket.accept()
    token = websocket.query_params.get("token")
    if not token: await websocket.close(code=4001); return
    payload = decode_token(token)
    if not payload: await websocket.close(code=4001); return
    with get_db() as db:
        member = db.execute("SELECT 1 FROM conversation_members WHERE conversation_id=? AND user_id=?", (cid, payload["sub"])).fetchone()
    if not member: await websocket.close(code=4003); return
    if cid not in ws_connections: ws_connections[cid] = []
    ws_connections[cid].append(websocket)
    try:
        while True:
            data = await websocket.receive_text()
            try:
                msg = json.loads(data)
                if msg.get("type") == "typing":
                    for ws in ws_connections.get(cid, []):
                        if ws != websocket:
                            try: await ws.send_text(json.dumps({"type": "typing", "user_id": payload["sub"], "conversation_id": cid}))
                            except: pass
                elif msg.get("type") == "signal":
                    for ws in ws_connections.get(cid, []):
                        if ws != websocket:
                            try: await ws.send_text(json.dumps(msg))
                            except: pass
            except: pass
    except WebSocketDisconnect:
        ws_connections[cid].remove(websocket)
        if not ws_connections[cid]: del ws_connections[cid]

# ═══════════════════════════════════════════
# CALLS (WebRTC Signaling)
# ═══════════════════════════════════════════
@app.websocket("/api/calls/ws/{cid}")
async def call_ws(websocket: WebSocket, cid: int):
    await websocket.accept()
    token = websocket.query_params.get("token")
    payload = decode_token(token) if token else None
    if not payload: await websocket.close(code=4001); return
    user_id = payload["sub"]
    if cid not in call_signals: call_signals[cid] = []
    call_signals[cid].append(websocket)
    try:
        while True:
            data = await websocket.receive_text()
            for ws in call_signals.get(cid, []):
                if ws != websocket:
                    try: await ws.send_text(data)
                    except: pass
    except WebSocketDisconnect:
        call_signals[cid].remove(websocket)
        if not call_signals[cid]: del call_signals[cid]

@app.post("/api/calls/start")
def start_call(data: dict, request: Request):
    u = get_user(request)
    with get_db() as db:
        db.execute("INSERT INTO calls (conversation_id, caller_id, call_type, status, started_at) VALUES (?,?,?,?,?)",
                   (data["conversation_id"], u["id"], data.get("call_type","voice"), "ringing", now_iso()))
        cid = db.execute("SELECT last_insert_rowid()").fetchone()[0]
        if data["conversation_id"] in ws_connections:
            for ws in ws_connections[data["conversation_id"]]:
                try:
                    asyncio.get_event_loop().create_task(ws.send_text(json.dumps({
                        "type": "call_incoming", "call_id": cid, "caller_id": u["id"],
                        "caller_name": u.get("full_name", u["username"]),
                        "call_type": data.get("call_type", "voice"),
                        "conversation_id": data["conversation_id"]
                    })))
                except: pass
        return {"call_id": cid}

@app.post("/api/calls/end")
def end_call(data: dict, request: Request):
    u = get_user(request)
    with get_db() as db:
        db.execute("UPDATE calls SET status='ended', ended_at=?, duration_seconds=? WHERE id=?",
                   (now_iso(), data.get("duration", 0), data["call_id"]))
    return {"ok": True}

# ═══════════════════════════════════════════
# FILES
# ═══════════════════════════════════════════
@app.post("/api/transfer/upload")
async def upload_file(request: Request, file: UploadFile = File(...), message: str = Form(""), is_public: bool = Form(False)):
    u = get_user(request)
    unique = f"{secrets.token_urlsafe(16)}_{file.filename}"
    content = await file.read()
    with open(os.path.join(UPLOAD_DIR, unique), "wb") as f: f.write(content)
    token = secrets.token_urlsafe(32) if is_public else None
    with get_db() as db:
        db.execute("INSERT INTO data_transfers (filename,original_name,file_size,mime_type,sender_id,message,is_public,share_token,created_at) VALUES (?,?,?,?,?,?,?,?,?)",
                   (unique, file.filename, len(content), file.content_type, u["id"], message, int(is_public), token, now_iso()))
        fid = db.execute("SELECT last_insert_rowid()").fetchone()[0]
    return {"ok": True, "id": fid, "filename": file.filename, "share_token": token}

@app.get("/api/transfer/files")
def list_files(request: Request):
    u = get_user(request)
    with get_db() as db:
        return rows_to_list(db.execute("SELECT * FROM data_transfers WHERE sender_id=? OR receiver_id=? ORDER BY created_at DESC", (u["id"], u["id"])).fetchall())

@app.get("/api/transfer/download/{tid}")
def download_file(tid: int, request: Request):
    u = get_user(request)
    with get_db() as db:
        t = db.execute("SELECT * FROM data_transfers WHERE id=?", (tid,)).fetchone()
        if not t: raise HTTPException(404, "Not found")
        db.execute("UPDATE data_transfers SET download_count=download_count+1 WHERE id=?", (tid,))
        return FileResponse(os.path.join(UPLOAD_DIR, t["filename"]), filename=t["original_name"], media_type=t["mime_type"] or "application/octet-stream")

@app.get("/api/transfer/share/{token}")
def download_shared(token: str):
    with get_db() as db:
        t = db.execute("SELECT * FROM data_transfers WHERE share_token=? AND is_public=1", (token,)).fetchone()
        if not t: raise HTTPException(404, "Not found")
        db.execute("UPDATE data_transfers SET download_count=download_count+1 WHERE id=?", (t["id"],))
        return FileResponse(os.path.join(UPLOAD_DIR, t["filename"]), filename=t["original_name"])

@app.delete("/api/transfer/files/{tid}")
def delete_file(tid: int, request: Request):
    u = get_user(request)
    with get_db() as db:
        t = db.execute("SELECT * FROM data_transfers WHERE id=? AND sender_id=?", (tid, u["id"])).fetchone()
        if not t: raise HTTPException(404, "Not found")
        path = os.path.join(UPLOAD_DIR, t["filename"])
        if os.path.exists(path): os.remove(path)
        db.execute("DELETE FROM data_transfers WHERE id=?", (tid,))
    return {"ok": True}

@app.get("/api/transfer/stats")
def transfer_stats(request: Request):
    u = get_user(request)
    with get_db() as db:
        s = db.execute("SELECT COUNT(*) c, COALESCE(SUM(file_size),0) s FROM data_transfers WHERE sender_id=?", (u["id"],)).fetchone()
        r = db.execute("SELECT COUNT(*) c, COALESCE(SUM(file_size),0) s FROM data_transfers WHERE receiver_id=?", (u["id"],)).fetchone()
        return {"sent_count": s["c"], "sent_size": s["s"], "received_count": r["c"], "received_size": r["s"]}

# ═══════════════════════════════════════════
# FRONTEND
# ═══════════════════════════════════════════
@app.get("/", response_class=HTMLResponse)
def index(request: Request):
    token = request.cookies.get("token")
    if not token or not decode_token(token): return RedirectResponse("/login")
    return templates.TemplateResponse(request, "index.html")

@app.get("/login", response_class=HTMLResponse)
def login_page(request: Request):
    return templates.TemplateResponse(request, "login.html")

@app.get("/invite/{token}")
def invite_page(request: Request, token: str):
    return RedirectResponse(f"/login?token={token}")

@app.get("/chat/join/{token}", response_class=HTMLResponse)
def join_chat_page(request: Request, token: str):
    return templates.TemplateResponse(request, "index.html")

if __name__ == "__main__":
    import uvicorn
    print("\n  PowerApp v2.0")
    print("  http://localhost:8000")
    print("  SuperAdmin: superadmin / superadmin\n")
    uvicorn.run(app, host="0.0.0.0", port=8000)
